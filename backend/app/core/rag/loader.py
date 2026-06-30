from pathlib import Path
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
from bs4 import BeautifulSoup
import markdown


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.html', '.pdf', '.docx', '.xlsx'}

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.splitter = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )

    def load_file(self, file_path: str) -> list[Document]:
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")

        text = self._extract_text(path, extension)

        doc = Document(
            text=text,
            metadata={
                "filename": path.name,
                "file_type": extension[1:],
            }
        )

        return [doc]

    def _extract_text(self, path: Path, extension: str) -> str:
        if extension == '.txt':
            return path.read_text(encoding='utf-8')

        elif extension == '.md':
            md_text = path.read_text(encoding='utf-8')
            html = markdown.markdown(md_text)
            return BeautifulSoup(html, 'html.parser').get_text()

        elif extension == '.html':
            html = path.read_text(encoding='utf-8')
            return BeautifulSoup(html, 'html.parser').get_text()

        elif extension == '.pdf':
            return self._extract_pdf(path)

        elif extension == '.docx':
            return self._extract_docx(path)

        elif extension == '.xlsx':
            return self._extract_xlsx(path)

        return ""

    def _extract_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            text_parts.append(f"[Page {i+1}]\n{page_text}")
        return "\n\n".join(text_parts)

    def _extract_docx(self, path: Path) -> str:
        doc = DocxDocument(str(path))
        return "\n\n".join(para.text for para in doc.paragraphs if para.text)

    def _extract_xlsx(self, path: Path) -> str:
        wb = load_workbook(str(path), read_only=True)
        text_parts = []
        for sheet in wb.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text:
                    rows.append(row_text)
            if rows:
                text_parts.append(f"[Sheet: {sheet.title}]\n" + "\n".join(rows))
        return "\n\n".join(text_parts)

    def chunk_documents(self, documents: list[Document]) -> list:
        return self.splitter.get_nodes_from_documents(documents)
